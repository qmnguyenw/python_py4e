Python | Working with .docx module



Word documents contain formatted text wrapped within three object levels.
Lowest level- Run objects, Middle level- Paragraph objects and Highest level-
Document object.  
So, we cannot work with these documents using normal text editors. But, we can
manipulate these word documents in python using the python-docx module.

>  **1.** The first step is to install this third-party module python-docx.
> You can use pip “pip install python-docx” or download the tarball from here.
> Here’s the Github repository.
>
>  **2.** After installation import “docx” NOT “python-docx”.  
>  **3.** Use “docx.Document” class to start working with the word document.

 **Code #1:**

 __

 __  
 __

 __

 __  
 __  
 __

# import docx NOT python-docx

import docx

 

# create an instance of a word document

doc = docx.Document()

 

# add a heading of level 0 (largest heading)

doc.add_heading('Heading for the document', 0)

 

# add a paragraph and store 

# the object in a variable

doc_para = doc.add_paragraph('Your paragraph goes here, ')

 

# add a run i.e, style like 

# bold, italic, underline, etc.

doc_para.add_run('hey there, bold here').bold = True

doc_para.add_run(', and ')

doc_para.add_run('these words are italic').italic = True

 

# add a page break to start a new page

doc.add_page_break()

 

# add a heading of level 2

doc.add_heading('Heading level 2', 2)

 

# pictures can also be added to our word document

# width is optional

doc.add_picture('path_to_picture')

 

# now save the document to a location

doc.save('path_to_document')  
  
---  
  
 __

 __

 **Output:**

  

  

![](https://media.geeksforgeeks.org/wp-content/uploads/word-document-
page-1-1.png)  
![](https://media.geeksforgeeks.org/wp-content/uploads/word-document-
page-2-1.png)  
Notice the page break in the second page.  
  
**Code #2:** Now, to open a word document, create an instance along with
passing the path to the document.

 __

 __  
 __

 __

 __  
 __  
 __

# import the Document class

# from the docx module

from docx import Document

 

# create an instance of a 

# word document we want to open

doc = Document('path_to_the_document')

 

# print the list of paragraphs in the document

print('List of paragraph objects:->>>')

print(doc.paragraphs)

 

# print the list of the runs 

# in a specified paragraph

print('\nList of runs objects in 1st paragraph:->>>')

print(doc.paragraphs[0].runs)

 

# print the text in a paragraph 

print('\nText in the 1st paragraph:->>>')

print(doc.paragraphs[0].text)

 

# for printing the complete document

print('\nThe whole content of the document:->>>\n')

for para in doc.paragraphs:

 print(para.text)  
  
---  
  
 __

 __

 **Output:**

    
    
    List of paragraph objects:->>>
    [<docx.text.paragraph.Paragraph object at 0x7f45b22dc128>,
    <docx.text.paragraph.Paragraph object at 0x7f45b22dc5c0>,
    <docx.text.paragraph.Paragraph object at 0x7f45b22dc0b8>,
    <docx.text.paragraph.Paragraph object at 0x7f45b22dc198>,
    <docx.text.paragraph.Paragraph object at 0x7f45b22dc0f0>]
    
    List of runs objects in 1st paragraph:->>>
    [<docx.text.run.Run object at 0x7f45b22dc198>]
    
    Text in the 1st paragraph:->>>
    Heading for the document
    
    The whole content of the document:->>>
    
    Heading for the document
    Your paragraph goes here, hey there, bold here, and these words are italic
    
    
    Heading level 2
    
    

**Reference:**https://python-docx.readthedocs.io/en/latest/#user-guide.

Attention reader! Don’t stop learning now. Get hold of all the important DSA
concepts with the **DSA Self Paced Course** at a student-friendly price and
become industry ready.

My Personal Notes _arrow_drop_up_

Save

